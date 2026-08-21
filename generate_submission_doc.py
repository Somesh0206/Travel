import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_document():
    doc = docx.Document()

    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading("HACKATHON ASSET SUBMISSION DOCUMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Team & General Info
    p = doc.add_paragraph()
    p.add_run("Team Name: ").bold = True
    p.add_run("MOVA\n")
    p.add_run("Project Name: ").bold = True
    p.add_run("MOVA - Accessible & Intelligent Public Transit Assistant\n")
    p.add_run("Live Production URL: ").bold = True
    p.add_run("https://travel-beige-iota.vercel.app\n")
    p.add_run("Primary GitHub Repository: ").bold = True
    p.add_run("https://github.com/Somesh0206/Travel.git\n")

    # Designated Trader Nomination
    doc.add_heading("1. Designated Trader Nomination (The Pit)", level=1)
    p_trader = doc.add_paragraph()
    p_trader.add_run("Nominated Trader Name: ").bold = True
    p_trader.add_run("Somesh Satapathy\n")
    p_trader.add_run("Role / Focus: ").bold = True
    p_trader.add_run("Lead Full-Stack & System Integration Engineer\n")
    p_trader.add_run("Contact Email / Discord: ").bold = True
    p_trader.add_run("someshsatapathy2007@gmail.com\n")

    # Listed Modular Assets (1 to 3 Assets)
    doc.add_heading("2. Modular Project Assets", level=1)

    # Asset 1
    doc.add_heading("Asset 1: MOVA-RoadRouter", level=2)
    p1 = doc.add_paragraph()
    p1.add_run("Asset Name: ").bold = True
    p1.add_run("MOVA-RoadRouter\n")
    p1.add_run("Category / Type: ").bold = True
    p1.add_run("Mapping, Geocoding & Road-Following Navigation Engine\n")
    p1.add_run("Description: ").bold = True
    p1.add_run(
        "A modular, standalone routing system that combines OpenStreetMap/OSRM road networks with "
        "interactive Leaflet visualization. It computes real road curves, turn-by-turn geometries, "
        "glowing multi-layer illuminated asphalt paths, dynamic map auto-framing, and simulated real-time "
        "vehicle transit animation along roads.\n"
    )
    p1.add_run("Standalone Source Code / Repository Link: ").bold = True
    p1.add_run("https://github.com/Somesh0206/Travel/blob/main/frontend/src/components/MapView.jsx\n")
    p1.add_run("Demo Video Link (30–45s): ").bold = True
    p1.add_run("https://youtu.be/demo-mova-roadrouter\n")
    p1.add_run("Key Technologies: ").bold = True
    p1.add_run("React, Leaflet, OSRM API, Turf.js, Web Animations\n")

    # Asset 2
    doc.add_heading("Asset 2: MOVA-SafetySOS", level=2)
    p2 = doc.add_paragraph()
    p2.add_run("Asset Name: ").bold = True
    p2.add_run("MOVA-SafetySOS\n")
    p2.add_run("Category / Type: ").bold = True
    p2.add_run("Emergency Alert Dispatcher & Audio Siren Deterrence Module\n")
    p2.add_run("Description: ").bold = True
    p2.add_run(
        "An autonomous emergency response module. Generates hardware-independent audible siren alarms "
        "using the Web Audio API (sawtooth oscillator) for attacker deterrence, and simultaneously captures "
        "live GPS coordinates to dispatch WhatsApp emergency broadcasts and query nearby police stations.\n"
    )
    p2.add_run("Standalone Source Code / Repository Link: ").bold = True
    p2.add_run("https://github.com/Somesh0206/Travel/blob/main/frontend/src/components/SOSButton.jsx\n")
    p2.add_run("Demo Video Link (30–45s): ").bold = True
    p2.add_run("https://youtu.be/demo-mova-safetysos\n")
    p2.add_run("Key Technologies: ").bold = True
    p2.add_run("Web Audio API, Geolocation API, WhatsApp Deep Linking, FastAPI\n")

    # Asset 3
    doc.add_heading("Asset 3: MOVA-OfflineTransitPack", level=3)
    p3 = doc.add_paragraph()
    p3.add_run("Asset Name: ").bold = True
    p3.add_run("MOVA-OfflineTransitPack\n")
    p3.add_run("Category / Type: ").bold = True
    p3.add_run("Zero-Connectivity Transit Timetable & Safety Cache Engine\n")
    p3.add_run("Description: ").bold = True
    p3.add_run(
        "A resilient offline transit engine allowing commuters in low/no-network areas to download complete "
        "bus schedules, campus gate locations, emergency hospital hotlines, and safety guides into local browser "
        "storage for instant offline search and navigation.\n"
    )
    p3.add_run("Standalone Source Code / Repository Link: ").bold = True
    p3.add_run("https://github.com/Somesh0206/Travel/blob/main/frontend/src/pages/Offline.jsx\n")
    p3.add_run("Demo Video Link (30–45s): ").bold = True
    p3.add_run("https://youtu.be/demo-mova-offlinepack\n")
    p3.add_run("Key Technologies: ").bold = True
    p3.add_run("LocalStorage API, JSON Caching, PWA-Ready Offline Sync\n")

    # Summary Table
    doc.add_heading("3. Asset Summary Table", level=1)
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["Asset ID & Name", "Core Functionality", "Repo / Gist Link", "Demo Link"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ("MOVA-RoadRouter", "Real-time OSRM road geometry & animated vehicle tracking", "https://github.com/Somesh0206/Travel/blob/main/frontend/src/components/MapView.jsx", "https://youtu.be/demo-mova-roadrouter"),
        ("MOVA-SafetySOS", "Web Audio API siren + WhatsApp emergency GPS dispatcher", "https://github.com/Somesh0206/Travel/blob/main/frontend/src/components/SOSButton.jsx", "https://youtu.be/demo-mova-safetysos"),
        ("MOVA-OfflineTransitPack", "Zero-connectivity timetable & emergency contact caching", "https://github.com/Somesh0206/Travel/blob/main/frontend/src/pages/Offline.jsx", "https://youtu.be/demo-mova-offlinepack")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            table.cell(row_idx, col_idx).text = text

    output_path = "c:\\Users\\KIIT\\Desktop\\Hackathon\\Demos\\Travel\\MOVA.docx"
    doc.save(output_path)
    print(f"Document created successfully at {output_path}")

if __name__ == "__main__":
    create_document()
